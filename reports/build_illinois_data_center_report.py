#!/usr/bin/env python3
from __future__ import annotations

import json
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "tmp/regulatory_workbook/regulatory_dataset.json"
OUT_DIR = ROOT / "reports"
WORK_DIR = ROOT / "tmp/report"
OUTPUT = OUT_DIR / "illinois-data-center-backup-power-report.docx"

NAVY = "17324D"
TEAL = "147D89"
TEAL_DARK = "0E5B63"
BLUE = "2E74B5"
GOLD = "C6922F"
INK = "1E2933"
MUTED = "66717C"
LIGHT_BLUE = "E8F2F5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RISK = "9B1C1C"
FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def pil_font(size: int, bold: bool = False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def draw_centered(draw, xy, text, font, fill, anchor="mm"):
    draw.text(xy, text, font=font, fill=fill, anchor=anchor)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, color=INK, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_bottom_border(paragraph, color="D4DCE3", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run_el = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    r_pr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")
    r_pr.append(sz)
    run_el.append(r_pr)
    text = OxmlElement("w:t")
    text.text = "1"
    run_el.append(text)
    field.append(run_el)
    paragraph._p.append(field)


def add_bullet(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, 11, INK, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, 11, INK)
    else:
        run = p.add_run(text)
        set_run_font(run, 11, INK)


def add_numbered(doc: Document, title: str, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(title)
    set_run_font(run, 11, NAVY, bold=True)
    run = p.add_run(text)
    set_run_font(run, 11, INK)


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = False
    run = p.add_run(text)
    set_run_font(run, 8.5, MUTED, italic=True)


def add_figure(doc: Document, image_path: Path, width: float, alt_text: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    inline = p.add_run().add_picture(str(image_path), width=Inches(width))._inline
    inline.docPr.set("descr", alt_text)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    r = cap.add_run(caption)
    set_run_font(r, 9, MUTED, italic=True)


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"{label}  ")
    set_run_font(r, 11, TEAL_DARK, bold=True)
    r = p.add_run(text)
    set_run_font(r, 11, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "ILLINOIS DATA CENTERS  |  REGULATORY SCREENING REPORT"
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, 8.5, MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    add_page_field(footer)


def build_charts(data: dict) -> tuple[Path, Path, Path, dict]:
    selected = [row for row in data["generators"] if row.get("selected_for_workbook")]
    facilities_by_agency = defaultdict(list)
    for facility in data["facilities"]:
        facilities_by_agency[facility["agency_id"]].append(facility)

    agency = defaultdict(lambda: {"units": 0, "capacity_kw": 0, "groups": 0})
    for row in selected:
        record = agency[row["agency_id"]]
        record["units"] += int(row.get("quantity") or 0)
        record["capacity_kw"] += int(row.get("quantity") or 0) * int(row.get("rated_kw_each") or 0)
        record["groups"] += 1

    def label_for(agency_id: str) -> str:
        if agency_id == "170002468817":
            return "Aligned Energy (ORD-01/02)"
        facilities = facilities_by_agency.get(agency_id, [])
        if not facilities:
            return agency_id
        label = facilities[0].get("site_name") or facilities[0].get("organization_name") or agency_id
        return {
            "Microsoft Hoffman Estates (CHI05) Data Center": "Microsoft Hoffman Estates (CHI05)",
            "Elk Grove Village (CHI10-11-12) Data Center": "Elk Grove Village (CHI10-12)",
            "Microsoft Corp-Chicago Data Center": "Microsoft Northlake",
            "NTT Global Data Centers CH LLC": "NTT Global Data Centers",
        }.get(label, label)

    ranked = sorted(
        ((label_for(key), value["capacity_kw"] / 1000, value["units"]) for key, value in agency.items()),
        key=lambda item: item[1],
        reverse=True,
    )
    total_capacity_mw = sum(item[1] for item in ranked)
    total_units = sum(item[2] for item in ranked)

    top = ranked[:10][::-1]
    top_path = WORK_DIR / "top_capacity.png"
    image = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(image)
    label_x, plot_x, plot_width = 35, 540, 790
    top_y, row_h, bar_h = 60, 84, 48
    max_value = max(value for _, value, _ in top)
    for tick in range(0, 251, 50):
        x = plot_x + int(tick / 250 * plot_width)
        draw.line((x, top_y - 15, x, top_y + row_h * 10 - 18), fill="#E1E6EA", width=2)
        draw.text((x, top_y + row_h * 10 - 5), str(tick), font=pil_font(24), fill="#66717C", anchor="mt")
    for index, (name, value, units) in enumerate(top):
        y = top_y + index * row_h
        draw.text((label_x, y + bar_h / 2), name, font=pil_font(25), fill="#273746", anchor="lm")
        width = int(value / max_value * plot_width)
        color = "#147D89" if value >= 100 else "#2E5D7B"
        draw.rounded_rectangle((plot_x, y, plot_x + width, y + bar_h), radius=6, fill=color)
        draw.text((plot_x + width + 18, y + bar_h / 2), f"{value:,.1f} MW | {units} units", font=pil_font(23), fill="#334155", anchor="lm")
    draw.text((plot_x + plot_width / 2, 955), "Selected nameplate capacity (MW)", font=pil_font(25), fill="#4B5563", anchor="mm")
    image.save(top_path, dpi=(220, 220))

    elk_ids = {
        row["facility_id"]
        for row in data["facilities"]
        if "Elk Grove" in (row.get("manifest_address") or "")
    }
    elk_units = 0
    elk_capacity_mw = 0.0
    for row in selected:
        ids = {value.strip() for value in (row.get("facility_ids") or "").split(";") if value.strip()}
        if ids & elk_ids:
            elk_units += int(row.get("quantity") or 0)
            elk_capacity_mw += int(row.get("quantity") or 0) * int(row.get("rated_kw_each") or 0) / 1000

    concentration_path = WORK_DIR / "elk_grove_concentration.png"
    image = Image.new("RGB", (1600, 620), "white")
    draw = ImageDraw.Draw(image)
    panels = [
        (70, 735, [elk_capacity_mw, total_capacity_mw - elk_capacity_mw], "Nameplate capacity", "MW"),
        (865, 1530, [elk_units, total_units - elk_units], "Generator units", "units"),
    ]
    labels = ["Elk Grove area", "Other listed areas"]
    for left, right, values, title, unit in panels:
        draw.text(((left + right) / 2, 45), title, font=pil_font(30, True), fill="#17324D", anchor="mm")
        baseline, chart_top = 510, 120
        max_value = max(values) * 1.18
        for idx, value in enumerate(values):
            center = left + (idx + 0.5) * (right - left) / 2
            bar_width = 190
            height = int(value / max_value * (baseline - chart_top))
            fill = "#147D89" if idx == 0 else "#B8C8D1"
            draw.rounded_rectangle((center - bar_width / 2, baseline - height, center + bar_width / 2, baseline), radius=8, fill=fill)
            share = value / sum(values)
            display = f"{value:,.1f} {unit}" if unit == "MW" else f"{value:,.0f} {unit}"
            draw.text((center, baseline - height - 18), f"{display} | {share:.0%}", font=pil_font(24, True), fill="#334155", anchor="ms")
            draw.text((center, baseline + 22), labels[idx], font=pil_font(23), fill="#4B5563", anchor="mt")
        draw.line((left, baseline, right, baseline), fill="#D6DEE3", width=2)
    image.save(concentration_path, dpi=(220, 220))

    years = Counter()
    for row in data["documents"]:
        if row.get("document_type") == "Air Permit - Final" and row.get("document_date"):
            years[int(row["document_date"][:4])] += 1
    year_range = list(range(min(years), max(years) + 1))
    values = [years.get(year, 0) for year in year_range]
    timeline_path = WORK_DIR / "permit_timeline.png"
    image = Image.new("RGB", (1600, 620), "white")
    draw = ImageDraw.Draw(image)
    left, right, top_y, baseline = 115, 1540, 70, 500
    max_count = max(values)
    for tick in range(0, max_count + 1, 2):
        y = baseline - int(tick / max_count * (baseline - top_y))
        draw.line((left, y, right, y), fill="#E1E6EA", width=2)
        draw.text((left - 20, y), str(tick), font=pil_font(20), fill="#66717C", anchor="rm")
    slot = (right - left) / len(year_range)
    for index, (year, value) in enumerate(zip(year_range, values)):
        x0 = left + index * slot + slot * 0.16
        x1 = left + (index + 1) * slot - slot * 0.16
        height = int(value / max_count * (baseline - top_y))
        fill = "#C6922F" if year >= 2024 else "#2E5D7B"
        draw.rectangle((x0, baseline - height, x1, baseline), fill=fill)
        if year % 2 == 0 or year >= 2024:
            draw.text(((x0 + x1) / 2, baseline + 18), str(year), font=pil_font(19), fill="#66717C", anchor="mt")
    draw.rounded_rectangle((945, 75, 1500, 145), radius=10, fill="#FFF7E6")
    draw.text((1222, 110), "22 documents from 2024 onward", font=pil_font(24, True), fill="#7A5A00", anchor="mm")
    image.save(timeline_path, dpi=(220, 220))

    known_units = sum(int(row.get("quantity") or 0) for row in selected if row.get("rated_kw_each") is not None)
    large_units = sum(int(row.get("quantity") or 0) for row in selected if (row.get("rated_kw_each") or 0) >= 2500)
    reviewed_groups = sum(row.get("review_status") == "Reviewed" for row in selected)
    top5_capacity = sum(item[1] for item in ranked[:5])
    latest_selected_year_by_agency: dict[str, int] = {}
    for row in selected:
        permit_date = row.get("permit_date") or ""
        if len(permit_date) >= 4 and permit_date[:4].isdigit():
            latest_selected_year_by_agency[row["agency_id"]] = max(
                latest_selected_year_by_agency.get(row["agency_id"], 0),
                int(permit_date[:4]),
            )
    metrics_out = {
        "total_units": total_units,
        "total_capacity_mw": total_capacity_mw,
        "agency_count": len(agency),
        "top5_share": top5_capacity / total_capacity_mw,
        "known_units": known_units,
        "missing_kw_units": total_units - known_units,
        "large_units": large_units,
        "elk_units": elk_units,
        "elk_capacity_mw": elk_capacity_mw,
        "elk_facility_count": len(elk_ids),
        "reviewed_groups": reviewed_groups,
        "selected_groups": len(selected),
        "air_docs": sum(years.values()),
        "air_docs_2024plus": sum(value for year, value in years.items() if year >= 2024),
        "selected_agencies_2024plus": sum(
            year >= 2024 for year in latest_selected_year_by_agency.values()
        ),
    }
    return top_path, concentration_path, timeline_path, metrics_out


def build_report() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    top_chart, concentration_chart, timeline_chart, m = build_charts(data)

    doc = Document()
    style_document(doc)

    # First page: memo-masthead opening under the standard_business_brief preset.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("SCREENING REPORT  |  AUGUST 2026")
    set_run_font(r, 10, TEAL_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Illinois Data Centers and the Hidden Diesel Backup Fleet")
    set_run_font(r, 24, NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("A rapid assessment of air permits, generator inventories, operating limits, and compliance records")
    set_run_font(r, 13.5, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Prepared for environmental nonprofit research and investigation planning  |  Data through May 21, 2026")
    set_run_font(r, 9.5, MUTED, italic=True)
    add_bottom_border(p, color="AFC4CE", size="10")

    add_callout(
        doc,
        "CENTRAL FINDING",
        f"The listed data centers collectively resemble a large, geographically concentrated diesel-generation fleet. The selected regulatory inventories describe about {m['total_units']:,} generators and at least {m['total_capacity_mw'] / 1000:.2f} GW of standby nameplate capacity, even though the units are regulated facility by facility as emergency equipment.",
    )

    kpi = doc.add_table(rows=1, cols=4)
    set_table_geometry(kpi, [2340, 2340, 2340, 2340])
    kpi_values = [
        (f"{m['total_units']:,}", "selected generators"),
        (f"{m['total_capacity_mw'] / 1000:.2f} GW", "known nameplate capacity"),
        (f"{m['agency_count']}", "unique regulatory records"),
        (f"{m['top5_share']:.0%}", "capacity in five largest records"),
    ]
    for index, (value, label) in enumerate(kpi_values):
        cell = kpi.cell(0, index)
        set_cell_shading(cell, NAVY if index == 0 else LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(value)
        set_run_font(r, 17, WHITE if index == 0 else NAVY, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, 8.5, WHITE if index == 0 else MUTED)

    doc.add_heading("Executive summary", level=1)
    add_bullet(doc, "Scale: The selected inventories document a fleet whose nameplate capacity is numerically comparable to a large power station, although it is not a continuously operating power source.", "Scale:")
    add_bullet(doc, f"Concentration: Eight listed facilities with Elk Grove-area addresses account for {m['elk_units']:,} units and {m['elk_capacity_mw']:,.1f} MW - roughly {m['elk_capacity_mw'] / m['total_capacity_mw']:.0%} of documented capacity.", "Concentration:")
    add_bullet(doc, f"Modern build-out: {m['air_docs_2024plus']} of {m['air_docs']} air-permit documents are dated 2024 or later, indicating that the regulatory picture is changing quickly.", "Modern build-out:")
    add_bullet(doc, "Accountability: Compliance agreements include alleged emissions-limit exceedances, late emissions reports, missed deviation reporting, and construction of generators before obtaining a permit.", "Accountability:")

    doc.add_page_break()

    doc.add_heading("1 | A large and concentrated standby fleet", level=1)
    p = doc.add_paragraph(
        f"The strongest signal is cumulative scale. Across {m['agency_count']} unique Document Explorer records with a selected generator inventory, the workbook identifies {m['total_units']:,} units and {m['total_capacity_mw']:,.0f} MW of known electrical nameplate capacity. The shared ORD-01/ORD-02 record is counted once in these totals. Five records contain about {m['top5_share']:.0%} of the capacity."
    )
    add_figure(
        doc,
        top_chart,
        6.25,
        "Horizontal bar chart showing the ten largest selected data-center generator fleets by nameplate capacity in megawatts.",
        "Figure 1. Ten largest selected regulatory inventories by standby-generator capacity.",
    )
    add_source(doc, "Source: Illinois Data Center Regulatory Inventory, Facilities and Generators sheets. Shared ORD-01/ORD-02 record de-duplicated.")

    p = doc.add_paragraph()
    r = p.add_run("Generator count alone can mislead. ")
    set_run_font(r, 11, NAVY, bold=True)
    r = p.add_run(
        f"Edged's selected inventory contains 66 relatively small units totaling about 45 MW, while CyrusOne's 90-unit inventory totals 210 MW. Among the {m['known_units']:,} units with a recovered kW rating, roughly {m['large_units'] / m['known_units']:.0%} are rated at 2,500 kW or more. Capacity, engine technology, operating hours, and location are therefore more useful screening variables than unit count by itself."
    )
    set_run_font(r, 11, INK)

    doc.add_heading("Geographic concentration", level=2)
    add_figure(
        doc,
        concentration_chart,
        6.15,
        "Two bar charts comparing selected generator capacity and generator count in the Elk Grove area with all other listed areas.",
        f"Figure 2. Elk Grove-area facilities contain about {m['elk_capacity_mw'] / m['total_capacity_mw']:.0%} of selected capacity and {m['elk_units'] / m['total_units']:.0%} of selected units.",
    )
    add_source(doc, "Source: Facility manifest addresses and selected generator records. 'Elk Grove area' includes addresses containing Elk Grove or Elk Grove Village.")

    p = doc.add_paragraph(
        "This clustering reframes the issue as a neighborhood-scale cumulative-impact question. Routine testing, outages, or grid emergencies could cause many engines in the same industrial corridor to operate within a short period. The workbook cannot estimate exposure by itself, but it identifies where demographic, school, health, ambient-air, and wind analyses would be most valuable."
    )

    doc.add_page_break()

    doc.add_heading("2 | A rapidly changing regulatory landscape", level=1)
    p = doc.add_paragraph(
        f"The document set spans 2007 through May 2026, but recent activity is substantial: {m['air_docs_2024plus']} of {m['air_docs']} final air-permit documents are dated 2024 or later. {m['selected_agencies_2024plus']} of the {m['agency_count']} selected regulatory inventories rely on permit evidence from 2024 or later. Any public-facing inventory should therefore be maintained as a time series rather than treated as a one-time census."
    )
    add_figure(
        doc,
        timeline_chart,
        6.25,
        "Bar chart of final air-permit documents by year from 2007 through 2026, with years 2024 through 2026 highlighted.",
        "Figure 3. Final air-permit documents by year in the downloaded corpus.",
    )
    add_source(doc, "Source: Documents sheet. Counts reflect documents in this 25-facility screening sample, not all Illinois data centers.")

    doc.add_heading("Operating limits change the pollution picture", level=2)
    p = doc.add_paragraph(
        "Nameplate capacity is not annual generation or annual emissions. Extracted permits include individual-generator runtime limits from approximately 180 to 1,200 hours per year, aggregate fleet-hour limits, and monthly or annual diesel-consumption limits. Actual emissions depend on permitted and actual operating hours, load, engine age, emission factors, testing schedules, emergency use, fuel, and control equipment."
    )

    doc.add_heading("Compliance records show multiple failure modes", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3000, 1320, 5040])
    headers = ["Allegation theme", "Records", "What the agreements describe"]
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index != 1 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, 9.5, NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    compliance_rows = [
        ("Late Annual Emissions Reports", "3", "Equinix, CyrusOne Aurora, and Ensono reporting allegations."),
        ("Emission-limit exceedances", "2", "CO and VOM allegations involving Equinix generators G7-G10."),
        ("Deviation reporting", "1", "Alleged failure to notify the Compliance Section within 30 days."),
        ("Construction before permitting", "1", "Stream allegedly constructed four generators before obtaining an air construction permit."),
    ]
    for label, count, description in compliance_rows:
        cells = table.add_row().cells
        for index, value in enumerate((label, count, description)):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index != 1 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run_font(r, 9.5, INK, bold=(index == 0))
    add_source(doc, "Source: Compliance sheet. These are allegations summarized from executed compliance commitment agreements, not independent findings by this report.")

    add_callout(
        doc,
        "WHY IMAGERY MATTERS",
        "The Stream agreement is a concrete example of the monitoring opportunity: satellite imagery or aerial photography could help identify new generator yards or construction activity and compare observed timing with permit issuance.",
        fill="FFF7E6",
    )

    heading = doc.add_heading("3 | Implications for nonprofit research", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph(
        "The most compelling public-interest framing is that cloud reliability has a physical, local, and diesel-intensive footprint. The regulatory system evaluates engines permit by permit, while residents experience the cumulative mix of nearby sources. This dataset provides a starting point for testing whether those two scales diverge."
    )

    add_numbered(doc, "Verify the largest fleets. ", "Manually review the latest permits for CyrusOne, Aligned Energy, Prime, CHI10-12, and Elk Grove Data Center LLC before publishing exact totals.")
    add_numbered(doc, "Build a construction timeline. ", "Use historical imagery to date generator-pad expansion and compare observed changes with construction-permit and operating-permit dates.")
    add_numbered(doc, "Measure cumulative exposure. ", "Combine facility footprints with schools, homes, environmental-justice indicators, existing pollution sources, ambient monitors, and prevailing winds.")
    add_numbered(doc, "Move from capacity to emissions. ", "Extract generator-specific emission factors, permitted hours, fuel limits, testing conditions, and control requirements; then obtain actual runtime or fuel-use records where possible.")

    doc.add_heading("Interpretation limits", level=2)
    caveats = [
        "This is a 25-facility screening sample, not a statewide census of all Illinois data centers.",
        f"Only {m['reviewed_groups']} of {m['selected_groups']} selected generator-group records are marked Reviewed; the others remain OCR-derived and require confirmation.",
        "Selected inventory means the latest extractable operating-permit inventory plus later construction additions, or the most recent extractable permit when no operating inventory was available.",
        f"{m['missing_kw_units']} selected units lack a recovered kW rating and therefore contribute zero to the reported capacity total; {m['total_capacity_mw'] / 1000:.2f} GW should be treated as a minimum known value.",
        "A permit's presence does not establish that every authorized unit was constructed, remains installed, or operates during a given event.",
        "ORD-01 and ORD-02 share one regulatory record; this report counts that record once, but the allocation of its equipment between the two addresses remains unresolved.",
    ]
    for caveat in caveats:
        add_bullet(doc, caveat)

    doc.add_heading("Method and sources", level=2)
    p = doc.add_paragraph(
        "The analysis is based on the Illinois Data Center Regulatory Inventory workbook assembled from Illinois EPA Document Explorer records. The corpus contains 112 downloaded PDFs: 73 final air permits, five compliance documents, seven NPDES permits, and 27 site-remediation documents. OCR was performed on the 78 air-permit and compliance PDFs, totaling 1,181 pages. The report uses the workbook's selected generator records and de-duplicates the shared ORD-01/ORD-02 agency record for aggregate calculations."
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Primary source portal: ")
    set_run_font(r, 9, MUTED, bold=True)
    r = p.add_run("https://webapps.illinois.gov/EPA/DocumentExplorer/")
    set_run_font(r, 9, BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Underlying source URLs, document IDs, hashes, OCR paths, page references, and evidence excerpts are retained in the regulatory workbook.")
    set_run_font(r, 9, MUTED, italic=True)

    doc.core_properties.title = "Illinois Data Centers and the Hidden Diesel Backup Fleet"
    doc.core_properties.subject = "Screening analysis of Illinois data-center air permits, generators, and compliance records"
    doc.core_properties.author = "Earth Genome"
    doc.core_properties.keywords = "Illinois, data centers, diesel generators, air permits, compliance, environmental justice"
    doc.core_properties.created = datetime(2026, 8, 25)
    doc.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "metrics": m}, indent=2))


if __name__ == "__main__":
    build_report()
